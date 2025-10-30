import pandas as pd
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from pathlib import Path
import re

# --- CONFIGURAÇÕES DO LAYOUT DE ETIQUETAS (A4, 2 COLUNAS) ---
LARGURA_ETIQUETA = Cm(9.9)
ALTURA_ETIQUETA = Cm(3.4)
LARGURA_ESPACO = Cm(0.3)

# Mapeamento de chaves: [Chave Padrão] -> [Variações encontradas nos arquivos]
# --- ATUALIZADO ---
# Adicionadas as variações 'CPF/CNPJ', 'Endereço' e 'IBGE'
# que estavam faltando, com base no seu exemplo.
MAP_CHAVES_PADRAO = {
    'Nome': ['Nome completo / Razão Social', 'Nome'],
    'CPF': ['CPF ou CNPJ', 'CPF/CNPJ', 'CPF'],
    'Endereço': ['Endereço completo', 'Endereço Completo', 'Endereço'],
    'Telefone': ['Telefone de contato', 'Telefone'],
    'Qtd Cartões': ['Qtd de cartões', 'Qtda Cartões', 'qtd cartões', 'qtda de cartões'],
    'IBGE': ['IBGE de atuação', 'IBGE'],
}
# --- FIM DA ATUALIZAÇÃO ---


# --- FUNÇÕES DE LEITURA DE ARQUIVOS ---

def extrair_valor(texto_linha, map_chaves):
    """
    Verifica se uma linha de texto contém uma das chaves mapeadas.
    Retorna (chave_padrao, valor) ou (chave_padrao, None) se só a chave for encontrada.
    """
    texto_limpo = texto_linha.strip().strip(',').strip('"').strip()
    
    for chave_padrao, variacoes in map_chaves.items():
        for variacao in variacoes:
            # --- CASO 1: "Chave: Valor" ---
            # Regex agora é menos gulosa e para no fim da linha
            match = re.match(rf'^{re.escape(variacao)}\s*:\s*(.*)$', texto_limpo, re.IGNORECASE)
            if match:
                valor = match.group(1).strip().strip('"')
                # Se o valor estiver vazio, pode ser o caso "Chave:" (Valor na próxima)
                if valor:
                    # --- CORREÇÃO ANTI-DUPLICAÇÃO ---
                    # Se o valor capturado contiver *outra* chave, removemos.
                    # Isso limpa casos como "Valor1 Nome: Valor2"
                    for _, outras_variacoes in map_chaves.items():
                        for outra_var in outras_variacoes:
                            # Procura por " Nome: " ou " CPF: " etc.
                            match_interno = re.search(rf'\s+{re.escape(outra_var)}\s*:', valor, re.IGNORECASE)
                            if match_interno:
                                # Se encontrou, corta o valor antes da chave interna
                                valor = valor[:match_interno.start()].strip()
                                break
                        else:
                            continue # Continua no loop externo
                        break # Sai do loop externo
                    
                    return chave_padrao, valor
                else:
                    return chave_padrao, None # Encontrou "Chave:", mas sem valor
            
            # --- CASO 2: "Chave" (sozinha, talvez com :) ---
            # Compara a linha inteira (ignorando :) com a variação
            if texto_limpo.rstrip(':').strip().lower() == variacao.lower():
                 return chave_padrao, None # Encontrou só a chave

    return None, None # Não encontrou

def ler_arquivo_xlsx(caminho_arquivo, map_chaves, logger):
    """Lê um único arquivo Excel (xlsx) e extrai os dados."""
    logger(f"Processando XLSX: {caminho_arquivo.name}")
    dados_etiquetas = []
    try:
        df = pd.read_excel(caminho_arquivo, engine='openpyxl', header=None)
        if 0 not in df.columns:
            logger(f"AVISO: Arquivo {caminho_arquivo.name} não possui a coluna de dados (índice 0).")
            return []

        lista_de_celulas = df[0].astype(str).str.strip().tolist()
        registro_atual = {}
        ultima_chave_encontrada = None

        for celula in lista_de_celulas:
            texto_teste = re.sub(r'\s+', ' ', celula).strip()
            if not texto_teste or texto_teste.lower() == 'nan':
                continue

            # (Lógica similar ao docx para lidar com Chave/Valor separados)
            chave_padrao, valor = extrair_valor(texto_teste, map_chaves)
            
            if chave_padrao:
                if chave_padrao == 'Nome' and registro_atual:
                    if len(registro_atual) >= 3:
                        dados_etiquetas.append(registro_atual.copy())
                    registro_atual = {}
                
                if valor:
                    registro_atual[chave_padrao] = valor
                    ultima_chave_encontrada = chave_padrao
                else:
                    ultima_chave_encontrada = chave_padrao
            
            else:
                # Não é chave, pode ser um valor ou continuação
                if ultima_chave_encontrada:
                    linha_continua = texto_teste.strip().strip(',').strip('"')
                    
                    # --- NOVA VERIFICAÇÃO (igual ao docx) ---
                    # Se a linha "continua" parece ser uma nova chave:valor
                    # (mesmo que não esteja no MAP), não é uma continuação.
                    # Ex: "Rua / Avenida: ..." ou "Número: ..."
                    if re.match(r'^[a-zA-Z\s/()]+:\s*.+', linha_continua):
                        ultima_chave_encontrada = None
                        continue # Ignora esta linha (não é continuação)
                    # --- FIM DA NOVA VERIFICAÇÃO ---
                    
                    if ultima_chave_encontrada == 'Endereço':
                        if 'Endereço' in registro_atual and registro_atual['Endereço']:
                            registro_atual['Endereço'] += f" {linha_continua}"
                        else:
                            registro_atual['Endereço'] = linha_continua
                    elif ultima_chave_encontrada not in registro_atual:
                         registro_atual[ultima_chave_encontrada] = linha_continua
                    
                    if ultima_chave_encontrada != 'Endereço':
                         ultima_chave_encontrada = None

        # Salva o último registro do arquivo
        if registro_atual and len(registro_atual) >= 3:
            dados_etiquetas.append(registro_atual.copy())

    except Exception as e:
        logger(f"ERRO ao ler {caminho_arquivo.name}: {e}")
    
    return dados_etiquetas

def ler_arquivo_docx(caminho_arquivo, map_chaves, logger):
    logger(f"Processando DOCX: {caminho_arquivo.name}")
    dados_etiquetas = []
    try:
        doc = Document(caminho_arquivo)
        registro_atual = {}
        ultima_chave_encontrada = None # Estado para rastrear a última chave

        def processar_linha_de_texto(texto_linha, registro_atual, dados_etiquetas):
            nonlocal ultima_chave_encontrada
            
            # Limpa espaços e quebras de linha
            texto_teste = re.sub(r'\s+', ' ', texto_linha).strip()
            
            # --- Lógica de Separação (Dados Destinatario) ---
            if "dados destinatario" in texto_teste.lower():
                if registro_atual and len(registro_atual) >= 3:
                    dados_etiquetas.append(registro_atual.copy())
                registro_atual.clear()
                ultima_chave_encontrada = None
                return # Processou, pode sair

            # --- Lógica de Linha Vazia ---
            if not texto_teste:
                return # Ignora linha vazia, não reseta ultima_chave

            # --- Lógica de Extração de Chave/Valor ---
            chave_padrao, valor = extrair_valor(texto_teste, map_chaves)
            
            if chave_padrao:
                # --- Caso 1: Encontramos uma Chave ---
                if valor:
                    # Caso 1a: Chave: Valor (na mesma linha)
                    registro_atual[chave_padrao] = valor
                    ultima_chave_encontrada = chave_padrao # Atualiza, mas não espera continuação
                else:
                    # Caso 1b: Só a Chave (Valor na próxima célula/linha)
                    ultima_chave_encontrada = chave_padrao
            
            else:
                # --- Caso 2: Não encontramos chave (provavelmente um valor ou continuação) ---
                if ultima_chave_encontrada:
                    linha_continua = texto_teste.strip().strip(',').strip('"')
                    
                    # --- NOVA VERIFICAÇÃO ---
                    # Se a linha "continua" parece ser uma nova chave:valor
                    # (mesmo que não esteja no MAP), não é uma continuação.
                    # Ex: "Rua / Avenida: ..." ou "Número: ..."
                    if re.match(r'^[a-zA-Z\s/()]+:\s*.+', linha_continua):
                        # Parece uma nova chave.
                        # Se a chave anterior era 'Endereço', paramos de concatenar.
                        if ultima_chave_encontrada == 'Endereço':
                            ultima_chave_encontrada = None # Para de concatenar
                        # Se for outra chave (ex: Nome), já teria sido resetada.
                        # Mas por via das dúvidas:
                        ultima_chave_encontrada = None
                        return # Ignora esta linha (não é continuação)
                    # --- FIM DA NOVA VERIFICAÇÃO ---

                    # Se for endereço, concatena
                    if ultima_chave_encontrada == 'Endereço':
                        if 'Endereço' in registro_atual and registro_atual['Endereço']:
                            registro_atual['Endereço'] += f" {linha_continua}"
                        else:
                            registro_atual['Endereço'] = linha_continua
                    
                    # Se for outra chave, e ela ainda não tiver valor, atribui
                    elif ultima_chave_encontrada not in registro_atual:
                         registro_atual[ultima_chave_encontrada] = linha_continua
                    
                    # Reseta a 'ultima_chave' depois de usada (exceto para Endereço)
                    if ultima_chave_encontrada != 'Endereço':
                         ultima_chave_encontrada = None


        # Itera sobre os elementos do corpo (parágrafos e tabelas)
        for block in doc.element.body:
            if block.tag.endswith('p'):
                # É um parágrafo
                p = Paragraph(block, doc)
                processar_linha_de_texto(p.text, registro_atual, dados_etiquetas)
            
            elif block.tag.endswith('tbl'):
                # É uma tabela
                table = Table(block, doc)
                for row in table.rows:
                    # --- MUDANÇA CRÍTICA ---
                    # Processa CADA CÉLULA individualmente
                    for cell in row.cells:
                        # Processamos cada parágrafo dentro da célula
                        for paragraph in cell.paragraphs:
                             processar_linha_de_texto(paragraph.text, registro_atual, dados_etiquetas)
                    
                    # --- IMPORTANTE ---
                    # Ao final de uma LINHA da tabela, resetamos a 'ultima_chave'
                    # para evitar que um 'Valor' na próxima linha se junte a uma 'Chave' da anterior.
                    if ultima_chave_encontrada != 'Endereço':
                        ultima_chave_encontrada = None


        # Salva o último registro do arquivo
        if registro_atual and len(registro_atual) >= 3:
            dados_etiquetas.append(registro_atual.copy())

    except Exception as e:
        logger(f"ERRO ao ler {caminho_arquivo.name}: {e}")

    return dados_etiquetas


def ler_pasta(caminho_pasta, map_chaves, logger):
    """Lê todos os arquivos .xlsx e .docx de uma pasta."""
    logger(f"Buscando arquivos em: {caminho_pasta}")
    dados_etiquetas = []
    
    if not caminho_pasta.is_dir():
        logger(f"ERRO: Diretório não encontrado: {caminho_pasta}")
        return []

    arquivos = list(caminho_pasta.glob('*.xlsx')) + list(caminho_pasta.glob('*.docx'))

    if not arquivos:
        logger("AVISO: Nenhum arquivo (.xlsx ou .docx) encontrado no diretório.")
        return []

    for arquivo in arquivos:
        if arquivo.suffix == '.xlsx':
            dados_etiquetas.extend(ler_arquivo_xlsx(arquivo, map_chaves, logger))
        elif arquivo.suffix == '.docx':
            dados_etiquetas.extend(ler_arquivo_docx(arquivo, map_chaves, logger))
            
    return dados_etiquetas

def processar_entradas(lista_de_caminhos, logger):
    """
    Processa uma lista de caminhos (podendo ser arquivos ou uma pasta).
    Retorna a lista total de dados de etiquetas, sem duplicatas.
    """
    dados_etiquetas_total = []
    
    if not lista_de_caminhos:
        logger("ERRO: Nenhum arquivo ou pasta de entrada selecionado.")
        return []

    for caminho_str in lista_de_caminhos:
        caminho = Path(caminho_str)
        
        if not caminho.exists():
            logger(f"ERRO: Caminho não encontrado: {caminho}")
            continue
            
        if caminho.is_dir():
            dados_etiquetas_total.extend(ler_pasta(caminho, MAP_CHAVES_PADRAO, logger))
        elif caminho.is_file():
            if caminho.suffix == '.xlsx':
                dados_etiquetas_total.extend(ler_arquivo_xlsx(caminho, MAP_CHAVES_PADRAO, logger))
            elif caminho.suffix == '.docx':
                dados_etiquetas_total.extend(ler_arquivo_docx(caminho, MAP_CHAVES_PADRAO, logger))
            else:
                logger(f"AVISO: Arquivo ignorado (tipo não suportado): {caminho.name}")
        
    logger(f"\nTotal de {len(dados_etiquetas_total)} registros coletados (bruto).")

    # --- LÓGICA PARA REMOVER DUPLICATAS ---
    dados_unicos = []
    registros_vistos = set()

    for registro in dados_etiquetas_total:
        # Cria uma representação "hashable" do registro
        # Usando as chaves principais para definir a unicidade
        # Usamos .get() para não dar erro se a chave não existir
        chave_unica = (
            registro.get('Nome', '').strip().lower(),
            registro.get('CPF', '').strip().lower(),
            registro.get('Endereço', '').strip().lower()
        )
        
        # Se o registro (baseado na chave) ainda não foi visto
        if chave_unica not in registros_vistos:
            dados_unicos.append(registro)
            registros_vistos.add(chave_unica)
            
    logger(f"Total de {len(dados_unicos)} registros únicos encontrados.")
    return dados_unicos


# --- FUNÇÕES DE GERAÇÃO DO WORD ---

def ajustar_layout_a4(document):
    """Ajusta o tamanho da página para A4 e define margens mínimas."""
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(0.5)
    section.right_margin = Cm(0.5)

def formatar_texto_etiqueta(dados):
    """
    Formata os dados de um registro (dicionário) em uma string para o rótulo.
    Apenas inclui campos que existem e têm valor.
    """
    linhas = [
        f"Nome: {dados.get('Nome', '')}",
        f"CPF/CNPJ: {dados.get('CPF', '')}",
        f"Endereço: {dados.get('Endereço', '')}",
        f"Telefone: {dados.get('Telefone', '')}",
        f"Qtd Cartões: {dados.get('Qtd Cartões', '')}",
        f"IBGE: {dados.get('IBGE', '')}"
    ]
    # Filtra linhas cujo valor (após o ': ') não está vazio
    linhas_validas = [linha for linha in linhas if linha.split(': ', 1)[-1].strip()]
    return "\n".join(linhas_validas)

def set_cell_width(cell, width_cm):
    """Define a largura da célula usando manipulação de XML para precisão."""
    cell.width = width_cm
    tcW = cell._tc.get_or_add_tcPr().get_or_add_tcW()
    tcW.set(qn("w:w"), str(int(width_cm.twips)))
    tcW.set(qn("w:type"), "dxa")

def gerar_documento_word(dados_etiquetas, nome_arquivo_saida, logger):
    """Cria o documento Word e preenche a tabela de etiquetas."""
    
    if not dados_etiquetas:
        logger("Nenhum dado para gerar o documento Word.")
        return

    try:
        document = Document()
        ajustar_layout_a4(document)
        
        num_colunas_tabela = 3
        num_etiquetas_por_linha = 2 
        num_etiquetas = len(dados_etiquetas)
        num_linhas = (num_etiquetas + num_etiquetas_por_linha - 1) // num_etiquetas_por_linha
        
        table = document.add_table(rows=num_linhas, cols=num_colunas_tabela)
        table.autofit = False
        
        table.columns[0].width = LARGURA_ETIQUETA
        table.columns[1].width = LARGURA_ESPACO
        table.columns[2].width = LARGURA_ETIQUETA
        
        for row_idx in range(num_linhas):
            row = table.rows[row_idx]
            set_cell_width(row.cells[0], LARGURA_ETIQUETA)
            set_cell_width(row.cells[1], LARGURA_ESPACO)
            set_cell_width(row.cells[2], LARGURA_ETIQUETA)

            row.height = ALTURA_ETIQUETA
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY 

        logger(f"Gerando tabela de {num_linhas} linhas x {num_colunas_tabela} colunas...")

        for i in range(num_etiquetas):
            dados = dados_etiquetas[i]
            
            row_idx = i // num_etiquetas_por_linha
            col_idx_etiqueta = 0 if (i % num_etiquetas_por_linha == 0) else 2
            
            cell = table.cell(row_idx, col_idx_etiqueta)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            
            texto_etiqueta = formatar_texto_etiqueta(dados)
            
            run = p.add_run(texto_etiqueta) 
            run.font.size = Pt(8) 
            run.font.name = 'Arial'

        document.save(nome_arquivo_saida)
        logger(f"\nSUCESSO: Documento Word gerado em: {nome_arquivo_saida}")

    except Exception as e:
        logger(f"\nERRO FATAL ao gerar o documento Word: {e}")

