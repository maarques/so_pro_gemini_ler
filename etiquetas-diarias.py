import re
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE 
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

# --- CONFIGURAÇÕES DO SCRIPT ---

# 1. ESPECIFIQUE O DIRETÓRIO DE ENTRADA
# Mude este caminho para o diretório onde seus arquivos estão localizados.
DIRETORIO_ARQUIVOS = Path('./planilhas_input')

# 2. ESPECIFIQUE O NOME DO ARQUIVO DE ENTRADA E SAÍDA
NOME_ARQUIVO_ENTRADA = 'ETIQUETAS DIARIAS.docx' # Alterado para DOCX
NOME_ARQUIVO_WORD = 'Etiquetas_Diarias_Formatado.docx' 

# --- CONFIGURAÇÕES DO LAYOUT DE ETIQUETAS (A4, 2 COLUNAS) ---

# Largura e Altura das Etiquetas (REQUISITO: 99mm x 34mm)
LARGURA_ETIQUETA = Cm(9.9)  # 9.9cm (99mm)
ALTURA_ETIQUETA = Cm(3.4)   # 3.4cm (34mm)
LARGURA_ESPACO = Cm(0.3)    # 0.3cm (2mm) de espaçamento central

# O cálculo da largura total para A4 (21.0cm):
# Margem Esquerda (0.8) + Etiqueta 1 (9.9) + Espaço Central (0.2) + Etiqueta 2 (9.9) + Margem Direita (0.2) = 21.0cm.

def ajustar_layout_a4(document):
    """Ajusta o tamanho da página para A4 e define margens mínimas."""
    section = document.sections[0]
    # Tamanho A4 (21cm x 29.7cm)
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    
    # Margens ajustadas para um layout de 2 colunas de 9.9cm com 0.2cm de espaçamento
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(0.8) # Margem Esquerda: 8mm
    section.right_margin = Cm(0.6) # Margem Direita: 6mm


def ler_tabela_existente_do_word(diretorio, nome_arquivo_entrada):
    """
    Lê a primeira tabela do arquivo DOCX e extrai o conteúdo de todas as células.
    Isso garante quebra de linha e conteúdo sejam preservados.
    """
    print(f"Buscando arquivo Word de entrada em: {diretorio.resolve() / nome_arquivo_entrada}")
    
    arquivo_word_path = diretorio / nome_arquivo_entrada
    if not arquivo_word_path.exists():
        print(f"ERRO: Arquivo '{nome_arquivo_entrada}' não encontrado no diretório.")
        return []

    etiquetas_texto = []
    
    try:
        # Carrega o documento DOCX
        document = Document(arquivo_word_path)
        
        if not document.tables:
            print("AVISO: Nenhuma tabela encontrada no documento de entrada.")
            return []
            
        # Pega a primeira tabela (índice 0)
        table = document.tables[0]
        
        # Itera sobre todas as células da tabela, linha por linha, coluna por coluna
        for row in table.rows:
            for cell in row.cells:
                # O texto da célula é limpo apenas de espaços em branco externos.
                # A formatação interna (quebras de linha, etc.) é preservada pelo cell.text
                etiqueta = cell.text.strip()
                
                # Adiciona apenas se houver conteúdo significativo (mais de 10 caracteres)
                if etiqueta and len(etiqueta) > 10:
                    # Remove caracteres de controle que causam o erro do lxml (mesmo que improváveis em DOCX limpo)
                    etiqueta_limpa = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', etiqueta)
                    etiquetas_texto.append(etiqueta_limpa)
                    
    except Exception as e:
        print(f"ERRO ao ler ou processar o arquivo {nome_arquivo_entrada}: {e}")
        return []

    print(f"\nTotal de {len(etiquetas_texto)} registros de célula extraídos para reformatação.")
    return etiquetas_texto


def formatar_texto_etiqueta(dados):
    """
    Função simples que retorna o texto.
    """
    return dados


def set_cell_width(cell, width_cm):
    """Define a largura da célula usando manipulação de XML para precisão."""
    cell.width = width_cm
    tcW = cell._tc.get_or_add_tcPr().get_or_add_tcW()
    tcW.set(qn("w:w"), str(int(width_cm.twips)))
    tcW.set(qn("w:type"), "dxa")


def gerar_documento_word(dados_etiquetas, nome_arquivo):
    """Cria o novo documento Word e preenche a tabela de etiquetas."""
    
    if not dados_etiquetas:
        print("Nenhum dado para gerar o documento Word.")
        return

    document = Document()
    
    # 1. Configura a página para A4 e Margens
    ajustar_layout_a4(document)
    
    # Adiciona a tabela com 3 colunas: [Etiqueta 1], [Espaçador], [Etiqueta 2]
    num_colunas_tabela = 3
    num_etiquetas_por_linha = 2 
    num_etiquetas_validas = len(dados_etiquetas)
    num_linhas = (num_etiquetas_validas + num_etiquetas_por_linha - 1) // num_etiquetas_por_linha
    
    if num_linhas == 0:
        print("Nenhum conteúdo válido encontrado nas células para gerar a tabela.")
        return

    table = document.add_table(rows=num_linhas, cols=num_colunas_tabela)
    table.autofit = False
    
    # Define a largura das 3 colunas (Usando a função de XML para maior precisão)
    table.columns[0].width = LARGURA_ETIQUETA
    table.columns[1].width = LARGURA_ESPACO
    table.columns[2].width = LARGURA_ETIQUETA
    
    # Itera sobre as linhas da tabela no novo documento e define o layout
    for row_idx in range(num_linhas):
        row = table.rows[row_idx]
        set_cell_width(row.cells[0], LARGURA_ETIQUETA)
        set_cell_width(row.cells[1], LARGURA_ESPACO)
        set_cell_width(row.cells[2], LARGURA_ETIQUETA)

        # Configura a altura da linha para SER EXATAMENTE 3.4 cm
        try:
            row.height = ALTURA_ETIQUETA
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY 
        except Exception as e:
            print(f"AVISO: Falha ao definir altura da linha por API na linha {row_idx} (erro: {e}).")
        
        # --- Configuração Vertical e Espaçador (AGORA INCONDICIONAL) ---
        # Garante o alinhamento vertical para as células de dados (0 e 2)
        row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        row.cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Configura e limpa a célula espaçadora (coluna 1)
        cell_spacer = row.cells[1]
        cell_spacer.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Limpa parágrafos antigos
        for p in cell_spacer.paragraphs:
            p._element.getparent().remove(p._element)
        cell_spacer.add_paragraph().text = '' # Adiciona um parágrafo vazio
            
    print(f"Gerando tabela de {num_linhas} linhas x {num_colunas_tabela} colunas com o conteúdo extraído...")

    # Variável para rastrear a posição na lista de dados de entrada
    etiqueta_index = 0
    
    for row_idx in range(num_linhas):
        row = table.rows[row_idx]
        
        for col_idx_etiqueta in [0, 2]: # Apenas colunas de etiqueta (0 e 2)
            if etiqueta_index < len(dados_etiquetas):
                dados_texto = dados_etiquetas[etiqueta_index]
                
                cell = row.cells[col_idx_etiqueta]
                
                # O alinhamento vertical já foi definido incondicionalmente acima.
                
                # Limpa parágrafos antigos
                for p in cell.paragraphs:
                    p._element.getparent().remove(p._element)

                # Cria novos parágrafos para cada linha da etiqueta
                linhas = dados_texto.split('\n')
                for i, linha in enumerate(linhas):
                    p = cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    
                    run = p.add_run(linha)
                    run.font.size = Pt(8) # Fonte 8pt para caber
                    
                etiqueta_index += 1
            
            # O bloco do espaçador (coluna 1) foi movido para o início do loop da linha.


    # Salva o documento
    document.save(nome_arquivo)
    print(f"\nSUCESSO: Documento Word gerado com sucesso em: {nome_arquivo}")


# Função auxiliar para definir bordas (mantida, mas não utilizada no loop principal)
def set_cell_border(cell, **kwargs):
    """
    Define a borda da célula. (Não utilizada, mas mantida para referência de estilo)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Cria ou busca a tag de bordas
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # Define cada borda
    for border_name, border_props in kwargs.items():
        border_element = OxmlElement(f'w:{border_name}')
        for key, value in border_props.items():
            border_element.set(qn(f'w:{key}'), str(value))
        
        # Remove a borda antiga, se existir
        for element in tcBorders.findall(qn(f'w:{border_name}')):
            tcBorders.remove(element)

        tcBorders.append(border_element)


if __name__ == '__main__':
    # 1. Tenta criar o diretório de entrada se não existir
    if not DIRETORIO_ARQUIVOS.exists():
        DIRETORIO_ARQUIVOS.mkdir(parents=True, exist_ok=True)
        print(f"Diretório '{DIRETORIO_ARQUIVOS.name}' criado. Coloque seu arquivo '{NOME_ARQUIVO_ENTRADA}' aqui.")
    
    # 2. Processa os dados do documento DOCX
    dados_etiquetas = ler_tabela_existente_do_word(DIRETORIO_ARQUIVOS, NOME_ARQUIVO_ENTRADA)

    # 3. Gera o documento Word com a formatação de etiqueta
    gerar_documento_word(dados_etiquetas, NOME_ARQUIVO_WORD)
