import pandas as pd
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from pathlib import Path
import re
import typing 

# --- Constantes de Configuração ---

LARGURA_ETIQUETA = Cm(9.9)
ALTURA_ETIQUETA = Cm(3.4)
LARGURA_ESPACO = Cm(0.3)



# --- Classe de Extração de Dados ---

class ExtratorDados:
    """
    Responsável por ler e extrair dados de arquivos XLSX e DOCX.
    """
    def __init__(self, map_chaves: dict, logger: typing.Callable):
        self.map_chaves = map_chaves
        self.logger = logger

    def _extrair_valor(self, texto_linha: str) -> typing.Tuple[typing.Optional[str], typing.Optional[str]]:
        """
        Função auxiliar interna para extrair a chave e o valor de uma linha.
        """
        texto_limpo = texto_linha.strip().strip(',').strip('"').strip()
        
        for chave_padrao, variacoes in self.map_chaves.items():
            for variacao in variacoes:
                match = re.match(rf'^{re.escape(variacao)}\s*:\s*(.*)$', texto_limpo, re.IGNORECASE)
                if match:
                    valor = match.group(1).strip().strip('"')
                    if valor:
                        for _, outras_variacoes in self.map_chaves.items():
                            for outra_var in outras_variacoes:
                                match_interno = re.search(rf'\s+{re.escape(outra_var)}\s*:', valor, re.IGNORECASE)
                                if match_interno:
                                    valor = valor[:match_interno.start()].strip()
                                    break
                            else:
                                continue 
                            break 
                        
                        return chave_padrao, valor
                    else:
                        return chave_padrao, None # Encontrou "Chave:", mas sem valor
                
                if texto_limpo.rstrip(':').strip().lower() == variacao.lower():
                    return chave_padrao, None # Encontrou só a chave

        return None, None # Não encontrou

    def ler_arquivo_xlsx(self, caminho_arquivo: Path) -> typing.List[dict]:
        """
        Lê um arquivo .xlsx e extrai os dados da etiqueta.
        """
        self.logger(f"Processando XLSX: {caminho_arquivo.name}")
        dados_etiquetas = []
        try:
            df = pd.read_excel(caminho_arquivo, engine='openpyxl', header=None)
            if 0 not in df.columns:
                self.logger(f"AVISO: Arquivo {caminho_arquivo.name} não possui a coluna de dados (índice 0).")
                return []

            lista_de_celulas = df[0].astype(str).str.strip().tolist()
            registro_atual = {}
            ultima_chave_encontrada = None

            for celula in lista_de_celulas:
                texto_teste = re.sub(r'\s+', ' ', celula).strip()
                if not texto_teste or texto_teste.lower() == 'nan':
                    continue

                chave_padrao, valor = self._extrair_valor(texto_teste)
                
                if chave_padrao:
                    if chave_padrao == 'Nome' and registro_atual:
                        if len(registro_atual) >= 3:
                            dados_etiquetas.append(registro_atual.copy())
                        registro_atual = {}
                    
                    if valor:
                        registro_atual[chave_padrao] = valor
                        ultima_chave_encontrada = chave_padrao
                    else:
                        ultima_chave_encontrada = chave_padrao
                
                else:
                    if ultima_chave_encontrada:
                        linha_continua = texto_teste.strip().strip(',').strip('"')
                        
                        if re.match(r'^[a-zA-Z\s/()]+:\s*.+', linha_continua):
                            ultima_chave_encontrada = None
                            continue 
                        
                        if ultima_chave_encontrada == 'Endereço':
                            if 'Endereço' in registro_atual and registro_atual['Endereço']:
                                registro_atual['Endereço'] += f" {linha_continua}"
                            else:
                                registro_atual['Endereço'] = linha_continua
                        elif ultima_chave_encontrada not in registro_atual:
                            registro_atual[ultima_chave_encontrada] = linha_continua
                        
                        if ultima_chave_encontrada != 'Endereço':
                            ultima_chave_encontrada = None

            if registro_atual and len(registro_atual) >= 3:
                dados_etiquetas.append(registro_atual.copy())

        except Exception as e:
            self.logger(f"ERRO ao ler {caminho_arquivo.name}: {e}")
        
        return dados_etiquetas

    def ler_arquivo_docx(self, caminho_arquivo: Path) -> typing.List[dict]:
        """
        Lê um arquivo .docx e extrai os dados da etiqueta.
        """
        self.logger(f"Processando DOCX: {caminho_arquivo.name}")
        dados_etiquetas = []
        try:
            doc = Document(caminho_arquivo)
            registro_atual = {}
            ultima_chave_encontrada = None 

            def processar_linha_de_texto(texto_linha: str, registro_atual: dict, dados_etiquetas: list):
                nonlocal ultima_chave_encontrada
                
                texto_teste = re.sub(r'\s+', ' ', texto_linha).strip()
                
                if "dados destinatario" in texto_teste.lower():
                    if registro_atual and len(registro_atual) >= 3:
                        dados_etiquetas.append(registro_atual.copy())
                    registro_atual.clear()
                    ultima_chave_encontrada = None
                    return 

                if not texto_teste:
                    return 

                # Usa o método _extrair_valor da instância da classe
                chave_padrao, valor = self._extrair_valor(texto_teste)
                
                if chave_padrao:
                    if valor:
                        registro_atual[chave_padrao] = valor
                        ultima_chave_encontrada = chave_padrao 
                    else:
                        ultima_chave_encontrada = chave_padrao
                
                else:
                    if ultima_chave_encontrada:
                        linha_continua = texto_teste.strip().strip(',').strip('"')
                        
                        if re.match(r'^[a-zA-Z\s/()]+:\s*.+', linha_continua):
                            if ultima_chave_encontrada == 'Endereço':
                                ultima_chave_encontrada = None 
                            ultima_chave_encontrada = None
                            return 

                        if ultima_chave_encontrada == 'Endereço':
                            if 'Endereço' in registro_atual and registro_atual['Endereço']:
                                registro_atual['Endereço'] += f" {linha_continua}"
                            else:
                                registro_atual['Endereço'] = linha_continua
                        
                        elif ultima_chave_encontrada not in registro_atual:
                            registro_atual[ultima_chave_encontrada] = linha_continua
                        
                        if ultima_chave_encontrada != 'Endereço':
                            ultima_chave_encontrada = None


            for block in doc.element.body:
                if block.tag.endswith('p'):
                    p = Paragraph(block, doc)
                    processar_linha_de_texto(p.text, registro_atual, dados_etiquetas)
                
                elif block.tag.endswith('tbl'):
                    table = Table(block, doc)
                    for row in table.rows:
                        for cell in row.cells:
                            for paragraph in cell.paragraphs:
                                processar_linha_de_texto(paragraph.text, registro_atual, dados_etiquetas)
                        
                        if ultima_chave_encontrada != 'Endereço':
                            ultima_chave_encontrada = None


            if registro_atual and len(registro_atual) >= 3:
                dados_etiquetas.append(registro_atual.copy())

        except Exception as e:
            self.logger(f"ERRO ao ler {caminho_arquivo.name}: {e}")

        return dados_etiquetas


    def ler_pasta(self, caminho_pasta: Path) -> typing.List[dict]:
        """
        Lê todos os arquivos .xlsx e .docx de uma pasta.
        """
        self.logger(f"Buscando arquivos em: {caminho_pasta}")
        dados_etiquetas = []
        
        if not caminho_pasta.is_dir():
            self.logger(f"ERRO: Diretório não encontrado: {caminho_pasta}")
            return []

        arquivos = list(caminho_pasta.glob('*.xlsx')) + list(caminho_pasta.glob('*.docx'))

        if not arquivos:
            self.logger("AVISO: Nenhum arquivo (.xlsx ou .docx) encontrado no diretório.")
            return []

        for arquivo in arquivos:
            if arquivo.suffix == '.xlsx':
                dados_etiquetas.extend(self.ler_arquivo_xlsx(arquivo))
            elif arquivo.suffix == '.docx':
                dados_etiquetas.extend(self.ler_arquivo_docx(arquivo))
                
        return dados_etiquetas

# --- Classe de Processamento ---

class ProcessadorEtiquetas:
    """
    Orquestra a extração e processa os dados (ex: deduplicação).
    """
    def __init__(self, extrator: ExtratorDados, logger: typing.Callable):
        self.extrator = extrator
        self.logger = logger

    def processar_e_deduplicar(self, lista_de_caminhos: typing.List[str]) -> typing.List[dict]:
        """
        Processa uma lista de arquivos/pastas, extrai dados e remove duplicados.
        """
        dados_etiquetas_total = []
        
        if not lista_de_caminhos:
            self.logger("ERRO: Nenhum arquivo ou pasta de entrada selecionado.")
            return []

        for caminho_str in lista_de_caminhos:
            caminho = Path(caminho_str)
            
            if not caminho.exists():
                self.logger(f"ERRO: Caminho não encontrado: {caminho}")
                continue
                
            if caminho.is_dir():
                dados_etiquetas_total.extend(self.extrator.ler_pasta(caminho))
            elif caminho.is_file():
                if caminho.suffix == '.xlsx':
                    dados_etiquetas_total.extend(self.extrator.ler_arquivo_xlsx(caminho))
                elif caminho.suffix == '.docx':
                    dados_etiquetas_total.extend(self.extrator.ler_arquivo_docx(caminho))
                else:
                    self.logger(f"AVISO: Arquivo ignorado (tipo não suportado): {caminho.name}")
        
        self.logger(f"\nTotal de {len(dados_etiquetas_total)} registros coletados (bruto).")

        # Lógica de Deduplicação
        dados_unicos = []
        registros_vistos = set()

        for registro in dados_etiquetas_total:
            chave_unica = (
                registro.get('Nome', '').strip().lower(),
                registro.get('CPF', '').strip().lower(),
                registro.get('Endereço', '').strip().lower()
            )
            
            if chave_unica not in registros_vistos:
                dados_unicos.append(registro)
                registros_vistos.add(chave_unica)
                
        self.logger(f"Total de {len(dados_unicos)} registros únicos encontrados.")
        return dados_unicos

# --- Classe de Geração de Documento ---

class GeradorEtiquetasWord:
    """
    Gera o documento .docx final com as etiquetas formatadas.
    """
    def __init__(self, logger: typing.Callable, largura_etiqueta: Cm, altura_etiqueta: Cm, largura_espaco: Cm):
        self.logger = logger
        self.largura_etiqueta = largura_etiqueta
        self.altura_etiqueta = altura_etiqueta
        self.largura_espaco = largura_espaco

    def _ajustar_layout_a4(self, document: Document):
        """Define o layout da página para A4 com margens específicas."""
        section = document.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(0.5)
        section.right_margin = Cm(0.5)

    def _formatar_texto_etiqueta(self, dados: dict) -> str:
        """Formata o dicionário de dados em uma string de texto para a etiqueta."""
        linhas = [
            f"Nome: {dados.get('Nome', '')}",
            f"CPF/CNPJ: {dados.get('CPF', '')}",
            f"Endereço: {dados.get('Endereço', '')}",
            f"Telefone: {dados.get('Telefone', '')}",
            f"Qtd Cartões: {dados.get('Qtd Cartões', '')}",
            f"IBGE: {dados.get('IBGE', '')}"
        ]
        linhas_validas = [linha for linha in linhas if linha.split(': ', 1)[-1].strip()]
        return "\n".join(linhas_validas)

    def _set_cell_width(self, cell: Table.cell, width_cm: Cm):
        """Define a largura exata de uma célula da tabela."""
        cell.width = width_cm
        tcW = cell._tc.get_or_add_tcPr().get_or_add_tcW()
        tcW.set(qn("w:w"), str(int(width_cm.twips)))
        tcW.set(qn("w:type"), "dxa")

    def gerar_documento(self, dados_etiquetas: typing.List[dict], nome_arquivo_saida: str):
        """
        Cria e salva o documento Word com as etiquetas.
        """
        if not dados_etiquetas:
            self.logger("Nenhum dado para gerar o documento Word.")
            return

        try:
            document = Document()
            self._ajustar_layout_a4(document)
            
            num_colunas_tabela = 3
            num_etiquetas_por_linha = 2 
            num_etiquetas = len(dados_etiquetas)
            num_linhas = (num_etiquetas + num_etiquetas_por_linha - 1) // num_etiquetas_por_linha
            
            table = document.add_table(rows=num_linhas, cols=num_colunas_tabela)
            table.autofit = False
            
            table.columns[0].width = self.largura_etiqueta
            table.columns[1].width = self.largura_espaco
            table.columns[2].width = self.largura_etiqueta
            
            for row_idx in range(num_linhas):
                row = table.rows[row_idx]
                self._set_cell_width(row.cells[0], self.largura_etiqueta)
                self._set_cell_width(row.cells[1], self.largura_espaco)
                self._set_cell_width(row.cells[2], self.largura_etiqueta)

                row.height = self.altura_etiqueta
                row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY 

            self.logger(f"Gerando tabela de {num_linhas} linhas x {num_colunas_tabela} colunas...")

            for i in range(num_etiquetas):
                dados = dados_etiquetas[i]
                
                row_idx = i // num_etiquetas_por_linha
                col_idx_etiqueta = 0 if (i % num_etiquetas_por_linha == 0) else 2
                
                cell = table.cell(row_idx, col_idx_etiqueta)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                
                texto_etiqueta = self._formatar_texto_etiqueta(dados)
                
                run = p.add_run(texto_etiqueta) 
                run.font.size = Pt(8) 
                run.font.name = 'Arial'

            document.save(nome_arquivo_saida)
            self.logger(f"\nSUCESSO: Documento Word gerado em: {nome_arquivo_saida}")

        except Exception as e:
            self.logger(f"\nERRO FATAL ao gerar o documento Word: {e}")
